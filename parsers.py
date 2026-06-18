import docx
import pdfplumber

def extract_text(file_path: str) -> str:
    """Reads PDF or Word files and unifies into a single text output, including Word tables."""
    text_content = []
    
    if file_path.endswith('.pdf'):
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text_content.append(page.extract_text() or "")
                
    elif file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        
        # 1. Extract regular paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())
        
        # 2. Extract text hidden inside Word Tables (Crucial for invoices!)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    # Join cell data with a space to help the LLM see the alignment
                    text_content.append(" | ".join(row_text))
                    
    else:
        raise ValueError(f"Unsupported file format provided: {file_path}")
        
    return "\n".join(text_content).strip()